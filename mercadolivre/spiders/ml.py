#IMPORTANTE
#//div/div/div[2]/div[2]/div[1]/div[2]/div[1]/p[@class="ui-promotions-pill ui-pb-highlight-wrapper coupon"]
import re
import requests
import unidecode
import scrapy
import requests
from docx import Document
import pandas
from datetime import datetime

start_row = 20  
end_row = 33
num_rows = end_row - start_row

db = pandas.read_excel("politica-promo.xlsx", engine='openpyxl')

db.columns = ['PRODUTO', 'SITE', 'COLUNA3', 'CLÁSSICO ML', 'COLUNA5', 'PREMIUM ML', 'COLUNA7', 'MARKETPLACES', 'COLUNA9']

df = pandas.read_excel("GESTÃO DE AÇÕES E-COMMERCE.xlsx", usecols='C:O', skiprows=start_row, nrows=num_rows, engine='openpyxl', sheet_name="POLÍTICA COMERCIAL Dez24")

df.columns = ['PRODUTO', 'inutil1', 'SITE', 'COLUNA3','inutil2', 'CLÁSSICO ML', 'COLUNA5','inutil3', 'PREMIUM ML', 'COLUNA7','inutil4', 'MARKETPLACES', 'COLUNA9']

for index, i in df.iterrows():
    if i['PRODUTO'] == "FONTE 40A":
        fonte40Marketplace = round(i['COLUNA3'], 2) - 0.02;
        fonte40Classico = round(i['COLUNA5'], 2) - 0.02;
        fonte40Premium = round(i['COLUNA7'], 2) - 0.02;
        fonte40PremiumPrice = round(i['PREMIUM ML'], 2) - 0.02;
        fonte40ClassicoPrice = round(i['CLÁSSICO ML'], 2) - 0.02;
        fonte40Marketplaceprice = round(i['SITE'], 2) - 0.02;
    elif i['PRODUTO'] == "FONTE 60A":
        fonte60Marketplace = round(i['COLUNA3'], 2) - 0.02;
        fonte60Classico = round(i['COLUNA5'], 2) - 0.02;
        fonte60Premium = round(i['COLUNA7'], 2) - 0.02;
        fonte60PremiumPrice = round(i['PREMIUM ML'], 2) - 0.02;
        fonte60ClassicoPrice = round(i['CLÁSSICO ML'], 2) - 0.02;
        fonte60Marketplaceprice = round(i['SITE'], 2) - 0.02;
    elif i['PRODUTO'] == "FONTE 60A LITE":
        fonte60liteMarketplace = round(i['COLUNA3'], 2) - 0.02;
        fonte60liteClassico = round(i['COLUNA5'], 2) - 0.02;
        fonte60litePremium = round(i['COLUNA7'], 2) - 0.02;
        fonte60litePremiumPrice = round(i['PREMIUM ML'], 2) - 0.02;
        fonte60liteClassicoPrice = round(i['CLÁSSICO ML'], 2) - 0.02;
        fonte60liteMarketplaceprice = round(i['SITE'], 2) - 0.02;
    elif i['PRODUTO'] == "FONTE 70A":
        fonte70Marketplace = round(i['COLUNA3'], 2) - 0.02;
        fonte70Classico = round(i['COLUNA5'], 2) - 0.02;
        fonte70Premium = round(i['COLUNA7'], 2) - 0.02;
        fonte70PremiumPrice = round(i['PREMIUM ML'], 2) - 0.02;
        fonte70ClassicoPrice = round(i['CLÁSSICO ML'], 2) - 0.02;
        fonte70Marketplaceprice = round(i['SITE'], 2) - 0.02;
    elif i['PRODUTO'] == "FONTE 70A LITE":
        fonte70liteMarketplace = round(i['COLUNA3'], 2) - 0.02;
        fonte70liteClassico = round(i['COLUNA5'], 2) - 0.02;
        fonte70litePremium = round(i['COLUNA7'], 2) - 0.02;
        fonte70litePremiumPrice = round(i['PREMIUM ML'], 2) - 0.02;
        fonte70liteClassicoPrice = round(i['CLÁSSICO ML'], 2) - 0.02;
        fonte70liteMarketplaceprice = round(i['SITE'], 2) - 0.02;
    elif i['PRODUTO'] == "FONTE 90 BOB":
        fonte90bobMarketplace = round(i['COLUNA3'], 2) - 0.02;
        fonte90bobClassico = round(i['COLUNA5'], 2) - 0.02;
        fonte90bobPremium = round(i['COLUNA7'], 2) - 0.02;
        fonte90bobPremiumPrice = round(i['PREMIUM ML'], 2) - 0.02;
        fonte90bobClassicoPrice = round(i['CLÁSSICO ML'], 2) - 0.02;
        fonte90bobMarketplaceprice = round(i['SITE'], 2) - 0.02;
    elif i['PRODUTO'] == "FONTE 120 BOB":
        fonte120bobMarketplace = round(i['COLUNA3'], 2) - 0.02;
        fonte120bobClassico = round(i['COLUNA5'], 2) - 0.02;
        fonte120bobPremium = round(i['COLUNA7'], 2) - 0.02;
        fonte120bobPremiumPrice = round(i['PREMIUM ML'], 2) - 0.02;
        fonte120bobClassicoPrice = round(i['CLÁSSICO ML'], 2) - 0.02;
        fonte120bobMarketplaceprice = round(i['SITE'], 2) - 0.02;
    elif i['PRODUTO'] == "FONTE 120A LITE":
        fonte120liteMarketplace = round(i['COLUNA3'], 2) - 0.02;
        fonte120liteClassico = round(i['COLUNA5'], 2) - 0.02;
        fonte120litePremium = round(i['COLUNA7'], 2) - 0.02;
        fonte120litePremiumPrice = round(i['PREMIUM ML'], 2) - 0.02;
        fonte120liteClassicoPrice = round(i['CLÁSSICO ML'], 2) - 0.02;
        fonte120liteMarketplaceprice = round(i['SITE'], 2) - 0.02;
    elif i['PRODUTO'] == "FONTE 120A":
        fonte120Marketplace = round(i['COLUNA3'], 2) - 0.02;
        fonte120Classico = round(i['COLUNA5'], 2) - 0.02;
        fonte120Premium = round(i['COLUNA7'], 2) - 0.02;
        fonte120PremiumPrice = round(i['PREMIUM ML'], 2) - 0.02;
        fonte120ClassicoPrice = round(i['CLÁSSICO ML'], 2) - 0.02;
        fonte120Marketplaceprice = round(i['SITE'], 2) - 0.02;
    elif i['PRODUTO'] == "FONTE 200 BOB":
        fonte200bobMarketplace = round(i['COLUNA3'], 2) - 0.02;
        fonte200bobClassico = round(i['COLUNA5'], 2) - 0.02;
        fonte200bobPremium = round(i['COLUNA7'], 2) - 0.02;
        fonte200bobPremiumPrice = round(i['PREMIUM ML'], 2) - 0.02;
        fonte200bobClassicoPrice = round(i['CLÁSSICO ML'], 2) - 0.02;
        fonte200bobMarketplaceprice = round(i['SITE'], 2) - 0.02;
    elif i['PRODUTO'] == "FONTE 200A LITE":
        fonte200liteMarketplace = round(i['COLUNA3'], 2) - 0.02;
        fonte200liteClassico = round(i['COLUNA5'], 2) - 0.02;
        fonte200litePremium = round(i['COLUNA7'], 2) - 0.02;
        fonte200litePremiumPrice = round(i['PREMIUM ML'], 2) - 0.02;
        fonte200liteClassicoPrice = round(i['CLÁSSICO ML'], 2) - 0.02;
        fonte200liteMarketplaceprice = round(i['SITE'], 2) - 0.02;
    elif i['PRODUTO'] == "FONTE 200 MONO":
        fonte200monoMarketplace = round(i['COLUNA3'], 2) - 0.02;
        fonte200monoClassico = round(i['COLUNA5'], 2) - 0.02;
        fonte200monoPremium = round(i['COLUNA7'], 2) - 0.02;
        fonte200monoPremiumPrice = round(i['PREMIUM ML'], 2) - 0.02;
        fonte200monoClassicoPrice = round(i['CLÁSSICO ML'], 2) - 0.02;
        fonte200monoMarketplaceprice = round(i['SITE'], 2) - 0.02;
    elif i['PRODUTO'] == "FONTE 200A":
        fonte200Marketplace = round(i['COLUNA3'], 2) - 0.02;
        fonte200Classico = round(i['COLUNA5'], 2) - 0.02;
        fonte200Premium = round(i['COLUNA7'], 2) - 0.02;
        fonte200PremiumPrice = round(i['PREMIUM ML'], 2) - 0.02;
        fonte200ClassicoPrice = round(i['CLÁSSICO ML'], 2) - 0.02;
        fonte200Marketplaceprice = round(i['SITE'], 2) - 0.02;
        
# for index, i in db.iterrows():
#     if i['PRODUTO'] == "FONTE 40A":
#         fonte40Marketplace = round(i['COLUNA3'], 2) - 0.02;
#         fonte40Classico = round(i['COLUNA5'], 2) - 0.02;
#         fonte40Premium = round(i['COLUNA7'], 2) - 0.02;
#         fonte40PremiumPrice = round(i['PREMIUM ML'], 2) - 0.02;
#         fonte40ClassicoPrice = round(i['CLÁSSICO ML'], 2) - 0.02;
#         fonte40Marketplaceprice = round(i['SITE'], 2) - 0.02;
#     elif i['PRODUTO'] == "FONTE 60A":
#         fonte60Marketplace = round(i['COLUNA3'], 2) - 0.02;
#         fonte60Classico = round(i['COLUNA5'], 2) - 0.02;
#         fonte60Premium = round(i['COLUNA7'], 2) - 0.02;
#         fonte60PremiumPrice = round(i['PREMIUM ML'], 2) - 0.02;
#         fonte60ClassicoPrice = round(i['CLÁSSICO ML'], 2) - 0.02;
#         fonte60Marketplaceprice = round(i['SITE'], 2) - 0.02;
#     elif i['PRODUTO'] == "FONTE 60A LITE":
#         fonte60liteMarketplace = round(i['COLUNA3'], 2) - 0.02;
#         fonte60liteClassico = round(i['COLUNA5'], 2) - 0.02;
#         fonte60litePremium = round(i['COLUNA7'], 2) - 0.02;
#         fonte60litePremiumPrice = round(i['PREMIUM ML'], 2) - 0.02;
#         fonte60liteClassicoPrice = round(i['CLÁSSICO ML'], 2) - 0.02;
#         fonte60liteMarketplaceprice = round(i['SITE'], 2) - 0.02;
#     elif i['PRODUTO'] == "FONTE 70A":
#         fonte70Marketplace = round(i['COLUNA3'], 2) - 0.02;
#         fonte70Classico = round(i['COLUNA5'], 2) - 0.02;
#         fonte70Premium = round(i['COLUNA7'], 2) - 0.02;
#         fonte70PremiumPrice = round(i['PREMIUM ML'], 2) - 0.02;
#         fonte70ClassicoPrice = round(i['CLÁSSICO ML'], 2) - 0.02;
#         fonte70Marketplaceprice = round(i['SITE'], 2) - 0.02;
#     elif i['PRODUTO'] == "FONTE 70A LITE":
#         fonte70liteMarketplace = round(i['COLUNA3'], 2) - 0.02;
#         fonte70liteClassico = round(i['COLUNA5'], 2) - 0.02;
#         fonte70litePremium = round(i['COLUNA7'], 2) - 0.02;
#         fonte70litePremiumPrice = round(i['PREMIUM ML'], 2) - 0.02;
#         fonte70liteClassicoPrice = round(i['CLÁSSICO ML'], 2) - 0.02;
#         fonte70liteMarketplaceprice = round(i['SITE'], 2) - 0.02;
#     elif i['PRODUTO'] == "FONTE 90 BOB":
#         fonte90bobMarketplace = round(i['COLUNA3'], 2) - 0.02;
#         fonte90bobClassico = round(i['COLUNA5'], 2) - 0.02;
#         fonte90bobPremium = round(i['COLUNA7'], 2) - 0.02;
#         fonte90bobPremiumPrice = round(i['PREMIUM ML'], 2) - 0.02;
#         fonte90bobClassicoPrice = round(i['CLÁSSICO ML'], 2) - 0.02;
#         fonte90bobMarketplaceprice = round(i['SITE'], 2) - 0.02;
#     elif i['PRODUTO'] == "FONTE 120 BOB":
#         fonte120bobMarketplace = round(i['COLUNA3'], 2) - 0.02;
#         fonte120bobClassico = round(i['COLUNA5'], 2) - 0.02;
#         fonte120bobPremium = round(i['COLUNA7'], 2) - 0.02;
#         fonte120bobPremiumPrice = round(i['PREMIUM ML'], 2) - 0.02;
#         fonte120bobClassicoPrice = round(i['CLÁSSICO ML'], 2) - 0.02;
#         fonte120bobMarketplaceprice = round(i['SITE'], 2) - 0.02;
#     elif i['PRODUTO'] == "FONTE 120A LITE":
#         fonte120liteMarketplace = round(i['COLUNA3'], 2) - 0.02;
#         fonte120liteClassico = round(i['COLUNA5'], 2) - 0.02;
#         fonte120litePremium = round(i['COLUNA7'], 2) - 0.02;
#         fonte120litePremiumPrice = round(i['PREMIUM ML'], 2) - 0.02;
#         fonte120liteClassicoPrice = round(i['CLÁSSICO ML'], 2) - 0.02;
#         fonte120liteMarketplaceprice = round(i['SITE'], 2) - 0.02;
#     elif i['PRODUTO'] == "FONTE 120A":
#         fonte120Marketplace = round(i['COLUNA3'], 2) - 0.02;
#         fonte120Classico = round(i['COLUNA5'], 2) - 0.02;
#         fonte120Premium = round(i['COLUNA7'], 2) - 0.02;
#         fonte120PremiumPrice = round(i['PREMIUM ML'], 2) - 0.02;
#         fonte120ClassicoPrice = round(i['CLÁSSICO ML'], 2) - 0.02;
#         fonte120Marketplaceprice = round(i['SITE'], 2) - 0.02;
#     elif i['PRODUTO'] == "FONTE 200 BOB":
#         fonte200bobMarketplace = round(i['COLUNA3'], 2) - 0.02;
#         fonte200bobClassico = round(i['COLUNA5'], 2) - 0.02;
#         fonte200bobPremium = round(i['COLUNA7'], 2) - 0.02;
#         fonte200bobPremiumPrice = round(i['PREMIUM ML'], 2) - 0.02;
#         fonte200bobClassicoPrice = round(i['CLÁSSICO ML'], 2) - 0.02;
#         fonte200bobMarketplaceprice = round(i['SITE'], 2) - 0.02;
#     elif i['PRODUTO'] == "FONTE 200A LITE":
#         fonte200liteMarketplace = round(i['COLUNA3'], 2) - 0.02;
#         fonte200liteClassico = round(i['COLUNA5'], 2) - 0.02;
#         fonte200litePremium = round(i['COLUNA7'], 2) - 0.02;
#         fonte200litePremiumPrice = round(i['PREMIUM ML'], 2) - 0.02;
#         fonte200liteClassicoPrice = round(i['CLÁSSICO ML'], 2) - 0.02;
#         fonte200liteMarketplaceprice = round(i['SITE'], 2) - 0.02;
#     elif i['PRODUTO'] == "FONTE 200 MONO":
#         fonte200monoMarketplace = round(i['COLUNA3'], 2) - 0.02;
#         fonte200monoClassico = round(i['COLUNA5'], 2) - 0.02;
#         fonte200monoPremium = round(i['COLUNA7'], 2) - 0.02;
#         fonte200monoPremiumPrice = round(i['PREMIUM ML'], 2) - 0.02;
#         fonte200monoClassicoPrice = round(i['CLÁSSICO ML'], 2) - 0.02;
#         fonte200monoMarketplaceprice = round(i['SITE'], 2) - 0.02;
#     elif i['PRODUTO'] == "FONTE 200A":
#         fonte200Marketplace = round(i['COLUNA3'], 2) - 0.02;
#         fonte200Classico = round(i['COLUNA5'], 2) - 0.02;
#         fonte200Premium = round(i['COLUNA7'], 2) - 0.02;
#         fonte200PremiumPrice = round(i['PREMIUM ML'], 2) - 0.02;
#         fonte200ClassicoPrice = round(i['CLÁSSICO ML'], 2) - 0.02;
#         fonte200Marketplaceprice = round(i['SITE'], 2) - 0.02;  

# if os.path.exists("dados_scrapy.docx"):
#     doc = Document("dados_scrapy.docx")
# else:

doc = Document()

def extract_price(response):
  price_selectors = [
      '//*[@id="price"]/div/div[1]/div[1]/span[1]/span/span[2]/text()',
      '//html/body/main/div[2]/div[5]/div/div[1]/div/div[1]/div/div[@class="ui-pdp-container__row ui-pdp-container__row--price"]/div/div[1]/div[1]/span/span/span[2]/text()',
      '//*[@id="ui-pdp-main-container"]/div[1]/div/div[1]/div[2]/div[3]/div[1]/div[1]/span/span/span[2]/text()',
      '//*[@id="ui-pdp-main-container"]/div[1]/div/div[1]/div[2]/div[2]/div[1]/div[1]/span[1]/span/span[2]/text()'
  ]
  
  for selector in price_selectors:
    price = response.xpath(selector).get()
    if price:
      price = price.replace('.', '')
      decimal_selector = selector.replace("span[2]/text()", "") + 'span[@class="andes-money-amount__cents andes-money-amount__cents--superscript-36"]/text()'
      price_decimal = response.xpath(decimal_selector).get()
      
      if price_decimal:
        return float(f"{price}.{price_decimal}")
      else:
        try:
          return float(price)
        except ValueError:
          pass

  return None  


def extract_price_new(response):
  price_selectors = [
      './/div/div/div[2]/div[2]/div[1]/div[1]/div/div/div/span[1]/span[@class="andes-money-amount__fraction"]/text()',
      './/div[1]/div[1]/div/div/div/span[1]/span[@class="andes-money-amount__fraction"]/text()'
  ]
  
  for selector in price_selectors:
    price = response.xpath(selector).get()
    if price:
      price = price.replace('.', '')
      decimal_selector = selector.replace('span[@class="andes-money-amount__fraction"]/text()', '') + 'span[@class="andes-money-amount__cents andes-money-amount__cents--superscript-24"]/text()'
      price_decimal = response.xpath(decimal_selector).get()
      
      if price_decimal:
        return float(f"{price}.{price_decimal}")
      else:
        try:
          return float(price)
        except ValueError:
          pass

  return None  


class MlSpider(scrapy.Spider):
    option_selected = ""
    option_selected_new = ""
    name = 'ml'
    start_urls = ["https://lista.mercadolivre.com.br/fonte-jfa"]
    
    def __init__(self, palavra=None, cookie=None, *args, **kwargs):
        super(MlSpider, self).__init__(*args, **kwargs)
        self.palavra = palavra
        self.cookie = cookie
    
    
    def parse(self, response, **kwargs):
        self.option_selected = self.palavra
        self.option_selected_new = self.palavra
        
        search_catalog = ""
        if self.option_selected == "FONTE 40A":
            search_catalog = "fonte 40a"
        if self.option_selected == "FONTE 60A LITE":
            search_catalog = "fonte 60a"
        elif self.option_selected == "FONTE 60A":
            search_catalog = "fonte 60a"
        if self.option_selected == "FONTE 70A LITE":
            search_catalog = "fonte 70a"
        elif self.option_selected == "FONTE 70A":
            search_catalog = "fonte 70a"
        elif self.option_selected == "FONTE 90 BOB":
            search_catalog = "fonte 90a"
        elif self.option_selected == "FONTE 120A":
            search_catalog = "fonte 120a"
        elif self.option_selected == "FONTE 120A LITE":
            search_catalog = "fonte 120a"
        elif self.option_selected == "FONTE 120 BOB":
            search_catalog = "fonte 120a"
        elif self.option_selected == "FONTE 200A":
            search_catalog = "fonte 200a"
        elif self.option_selected == "FONTE 200A LITE":
            search_catalog = "fonte 200a"
        elif self.option_selected == "FONTE 200 BOB":
            search_catalog = "fonte 200a"
        elif self.option_selected == "FONTE 200 MONO":
            search_catalog = "fonte 200a mono"
        search_catalog = search_catalog.replace(" ", "%20")
        sites = [ "https://www.bestonline.com.br", "https://www.shoppratico.com.br", "https://www.renovonline.com.br", "https://www.lsdistribuidora.com.br", "https://www.radicalsom.com.br"]
        if self.option_selected_new == "FONTE 40A":
            products = ['fonte-40a-jfa_OrderId_PRICE_NoIndex_True',
            ]
        elif self.option_selected_new == "FONTE 60A LITE":
            products = [
            'fonte-60a-lite-jfa_OrderId_PRICE_NoIndex_True',
            ]
        elif self.option_selected_new == "FONTE 60A":
            products = [
            'fonte-60a-jfa_OrderId_PRICE_NoIndex_True',
            ]
        elif self.option_selected_new == "FONTE 70A LITE":
            products = [
            'fonte-70a-lite-jfa_OrderId_PRICE_NoIndex_True',
            ]
        elif self.option_selected_new == "FONTE 70A":
            products = [
            'fonte-70a-jfa_OrderId_PRICE_NoIndex_True',
            ]
        elif self.option_selected_new == "FONTE 90 BOB":
            products = [
            'fonte-90a-bob-jfa_OrderId_PRICE_NoIndex_True',
            ]
        elif self.option_selected_new == "FONTE 120 BOB":
            products = [
            'fonte-120a-bob-jfa_OrderId_PRICE_NoIndex_True',
            ]
        elif self.option_selected_new == "FONTE 120A LITE":
            products = [
            'fonte-120a-lite-jfa_OrderId_PRICE_NoIndex_True',
            ]
        elif self.option_selected_new == "FONTE 120A":
            products = [
            'fonte-120a-jfa_OrderId_PRICE_NoIndex_True',
            ]
        elif self.option_selected_new == "FONTE 200 BOB":
            products = [
            'fonte-200a-bob-jfa_OrderId_PRICE_NoIndex_True',
            ]
        elif self.option_selected_new == "FONTE 200A LITE":
            products = [
            'fonte-200a-lite-jfa_OrderId_PRICE_NoIndex_True',
            ]
        elif self.option_selected_new == "FONTE 200 MONO":
            products = [
            'fonte-200a-mono-jfa_OrderId_PRICE_NoIndex_True'
            ]
        elif self.option_selected_new == "FONTE 200A":
            products = [
            'fonte-200a-jfa_OrderId_PRICE_NoIndex_True',
            ]
        # products = ['fonte-40a-jfa_OrderId_PRICE_NoIndex_True',
        #             'fonte-60a-jfa_OrderId_PRICE_NoIndex_True',
        #             'fonte-60a-lite-jfa_OrderId_PRICE_NoIndex_True',
        #             'fonte-70a-jfa_OrderId_PRICE_NoIndex_True',
        #             'fonte-70a-lite-jfa_OrderId_PRICE_NoIndex_True',
        #             'fonte-90a-bob-jfa_OrderId_PRICE_NoIndex_True',
        #             'fonte-120a-jfa_OrderId_PRICE_NoIndex_True',
        #             'fonte-120a-lite-jfa_OrderId_PRICE_NoIndex_True',
        #             'fonte-120a-bob-jfa_OrderId_PRICE_NoIndex_True',
        #             'fonte-200a-jfa_OrderId_PRICE_NoIndex_True',
        #             'fonte-200a-lite-jfa_OrderId_PRICE_NoIndex_True',
        #             'fonte-200a-bob-jfa_OrderId_PRICE_NoIndex_True',
        #             'fonte-200a-mono-jfa_OrderId_PRICE_NoIndex_True'
        #             ]
        
        for site in sites:
            for product in products:
                
                yield scrapy.Request(url=f'{site}/{product}', callback=self.parse_lojas)

                

    def finish(self, total_price, url, nomeFonte, loja, lugar):
        if self.option_selected_new == "FONTE 40A" and total_price >= fonte40Marketplace:
            return;
        elif self.option_selected_new == "FONTE 60A LITE" and total_price >= fonte60liteMarketplace:
            return;
        elif self.option_selected_new == "FONTE 60A" and total_price >= fonte60Marketplace:
            return;
        elif self.option_selected_new == "FONTE 70A LITE" and total_price >= fonte70liteMarketplace:
            return;
        elif self.option_selected_new == "FONTE 70A" and total_price >= fonte70Marketplace:
            return;
        elif self.option_selected_new == "FONTE 90 BOB" and total_price >= fonte90bobMarketplace:
            return;
        elif self.option_selected_new == "FONTE 120 BOB" and total_price >= fonte120bobMarketplace:
            return;
        elif self.option_selected_new == "FONTE 120A LITE" and total_price >= fonte120liteMarketplace:
            return;
        elif self.option_selected_new == "FONTE 120A" and total_price >= fonte120Marketplace:
            return;
        elif self.option_selected_new == "FONTE 200 BOB" and total_price >= fonte200bobMarketplace:
            return;
        elif self.option_selected_new == "FONTE 200A LITE" and total_price >= fonte200liteMarketplace:
            return;
        elif self.option_selected_new == "FONTE 200 MONO" and total_price >= fonte200monoMarketplace:
            return;
        elif self.option_selected_new == "FONTE 200A" and total_price >= fonte200Marketplace:
            return;
        
        
        parcelado = self.get_price_previsto("NA")

        doc.add_paragraph(f'Modelo: {self.option_selected_new}')
        doc.add_paragraph(f'URL: {url}')
        doc.add_paragraph(f'Nome: {nomeFonte}')
        doc.add_paragraph(f'Preço: {total_price}')
        doc.add_paragraph(f'Preço Previsto: {parcelado}')
        doc.add_paragraph(f'Loja: {loja}')
        doc.add_paragraph('Tipo: ')
        doc.add_paragraph(f'Lugar: {lugar}')
        doc.add_paragraph(f'Cupom: ')
        doc.add_paragraph("--------------------------------------------------------------------")
        doc.add_paragraph('')
        doc.save(fr"dados/{self.option_selected_new}.docx")
        if url != None:
            yield {
                'url': url,
                'name': nomeFonte,
                'price': total_price,
                'loja': loja,
                'tipo': "",
                'lugar': lugar
            }


    def parse_lojas(self, response):
        if "radicalsom" in response.url:
            loja = "RADICALSOM"
            lugar = "Artur nogueira, São Paulo."
        elif "bestonline" in response.url:
            loja = "LS DISTRIBUIDORA"
            lugar = "Elísio Medrado, Bahia"
        elif "shoppratico" in response.url:
            loja = "BESTONLINE"
            lugar = "Rosario, Santa Fe."
        elif "renovonline" in response.url:
            loja = "RENOV ONLINE"
            lugar = "São João da Boa Vista - SP"

        elif "lsdistribuidora" in response.url:
            loja = "LS Distribuidora"
            lugar = "Elísio Medrado - BR-BA"
        
        for i in response.xpath('//*[@id="root-app"]/div/div[3]/section/ol/li'):
            nomeFonte = i.xpath('.//a[@class="poly-component__title"]/text()').get()
            if nomeFonte is None:
                nomeFonte = i.xpath('.//a[@class="ui-search-link__title-card ui-search-link"]/text()').get()
            price = i.xpath('.//div/div/div[3]/div/div[1]/div/div/div/div/span/span[2]/text()').get()
            if price == None:
                price = i.xpath('.//div/div/div[2]/div/div[1]/div/div/span[@class="andes-money-amount andes-money-amount--cents-superscript"]/span[@class="andes-money-amount__fraction"]/text()').get()
            cents = i.xpath('.//div/div/div[3]/div/div[1]/div/div/div/div/span/span[4]/text()').get()
            if cents == None:
                cents = i.xpath('.//div/div/div[2]/div/div[1]/div/div/span[@class="andes-money-amount andes-money-amount--cents-superscript"]/span[@class="andes-money-amount__cents andes-money-amount__cents--superscript-24"]/text()').get()
            url = i.xpath('.//a[@class="poly-component__title"]/@href').get()
            if url is None:
                url = i.xpath('.//a[@class="ui-search-link__title-card ui-search-link"]/@href').get()
            
            
            nomeFonte = nomeFonte.lower()
            nomeFonte = unidecode.unidecode(nomeFonte)
            if not cents:
                cents = 0
            if price:
                price = price.replace('.', '')
                total_price = float(f"{price}.{cents}")
            if self.option_selected == "FONTE 40A":
                if "bob" not in nomeFonte and "lite" not in nomeFonte and "controle" not in nomeFonte and 'jfa' in nomeFonte:
                    if "40a" in nomeFonte or "40" in nomeFonte or "40 amperes" in nomeFonte or "40amperes" in nomeFonte or "36a" in nomeFonte or "36" in nomeFonte or "36 amperes" in nomeFonte or "36amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
                        return
                        
            elif self.option_selected == "FONTE 60A":
                if "bob" not in nomeFonte and "lite" not in nomeFonte and "controle" not in nomeFonte and 'jfa' in nomeFonte:
                    if "60a" in nomeFonte or "60" in nomeFonte or "60 amperes" in nomeFonte or "60amperes" in nomeFonte or "60 a" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
                        return
                        
            elif self.option_selected == "FONTE 60A LITE":
                if "bob" not in nomeFonte and "lite" in nomeFonte and "controle" not in nomeFonte and 'jfa' in nomeFonte:
                    if "60a" in nomeFonte or "60" in nomeFonte or "60 amperes" in nomeFonte or "60amperes" in nomeFonte or "60 a" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
                        return
                    
                        
            elif self.option_selected == "FONTE 70A":
                if "bob" not in nomeFonte and "lite" not in nomeFonte and "controle" not in nomeFonte and 'jfa' in nomeFonte:
                    if "70a" in nomeFonte or "70" in nomeFonte or "70 amperes" in nomeFonte or "70amperes" in nomeFonte or "70 a" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
                        return
                        
            elif self.option_selected == "FONTE 70A LITE":
                if "bob" not in nomeFonte and "lite" in nomeFonte and "controle" not in nomeFonte and 'jfa' in nomeFonte:
                    if "70a" in nomeFonte or "70" in nomeFonte or "70 amperes" in nomeFonte or "70amperes" in nomeFonte or "70 a" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
                        return
                    
                        
            elif self.option_selected == "FONTE 90 BOB":
                if "bob" in nomeFonte and "lite" not in nomeFonte and "controle" not in nomeFonte and 'jfa' in nomeFonte and 'mono' not in nomeFonte and 'monovolt' not in nomeFonte:
                    if "90a" in nomeFonte or "90" in nomeFonte or "90 amperes" in nomeFonte or "90amperes" in nomeFonte or "90 a" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
                        return
                    
                        
            elif self.option_selected == "FONTE 120A":
                if "bob" not in nomeFonte and "lite" not in nomeFonte and "controle" not in nomeFonte and 'jfa' in nomeFonte:
                    if "120a" in nomeFonte or "120" in nomeFonte or "120 amperes" in nomeFonte or "120amperes" in nomeFonte or "120 a" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
                        return
                        
            elif self.option_selected == "FONTE 120A LITE":
                if "bob" not in nomeFonte and "lite" in nomeFonte and "controle" not in nomeFonte and 'jfa' in nomeFonte:
                    if "120a" in nomeFonte or "120" in nomeFonte or "120 amperes" in nomeFonte or "120amperes" in nomeFonte or "120 a" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
                        return
                    
                        
            elif self.option_selected == "FONTE 120 BOB":
                if "bob" in nomeFonte and "lite" not in nomeFonte and "controle" not in nomeFonte and 'jfa' in nomeFonte and 'mono' not in nomeFonte and 'monovolt' not in nomeFonte:
                    if "120a" in nomeFonte or "120" in nomeFonte or "120 amperes" in nomeFonte or "120amperes" in nomeFonte or "120 a" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
                        return
                    
                        
            elif self.option_selected == "FONTE 200A":
                if "bob" not in nomeFonte and "lite" not in nomeFonte and "controle" not in nomeFonte and 'jfa' in nomeFonte and 'mono' not in nomeFonte and 'monovolt' not in nomeFonte and '220' not in nomeFonte and '220v' not in nomeFonte:
                    if "200a" in nomeFonte or "200" in nomeFonte or "200 amperes" in nomeFonte or "200amperes" in nomeFonte or "200 a" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
                        return
                        
            elif self.option_selected == "FONTE 200 MONO":
                if "bob" not in nomeFonte and "lite" not in nomeFonte and "controle" not in nomeFonte and 'jfa' in nomeFonte and ("mono" in nomeFonte or "220v" in nomeFonte or "monovolt" in nomeFonte):
                    if "200a" in nomeFonte or "200" in nomeFonte or "200 amperes" in nomeFonte or "200amperes" in nomeFonte or "200 a" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
                        return
                    
                        
            elif self.option_selected == "FONTE 200A LITE":
                if "bob" not in nomeFonte and "lite" in nomeFonte and "controle" not in nomeFonte and 'jfa' in nomeFonte and 'mono' not in nomeFonte and 'monovolt' not in nomeFonte:
                    if "200a" in nomeFonte or "200" in nomeFonte or "200 amperes" in nomeFonte or "200amperes" in nomeFonte or "200 a" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
                        return
                    
                        
            elif self.option_selected == "FONTE 200 BOB":
                if "bob" in nomeFonte and "lite" not in nomeFonte and "controle" not in nomeFonte and 'jfa' in nomeFonte and 'mono' not in nomeFonte and 'monovolt' not in nomeFonte:
                    if "200a" in nomeFonte or "200" in nomeFonte or "200 amperes" in nomeFonte or "200amperes" in nomeFonte or "200 a" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
                        return
                    
                       
        
    def get_price_previsto(self, tipo):
        if tipo == "Clássico":
            for index, i in db.iterrows():
                if self.option_selected_new == "FONTE 40A" and i['PRODUTO'] == "FONTE 40A":
                    return round(i['COLUNA5'], 2);
                elif self.option_selected_new == "FONTE 60A" and i['PRODUTO'] == "FONTE 60A":
                    return round(i['COLUNA5'], 2);
                elif self.option_selected_new == "FONTE 60A LITE" and i['PRODUTO'] == "FONTE 60A LITE":
                    return round(i['COLUNA5'], 2);
                elif self.option_selected_new == "FONTE 70A" and i['PRODUTO'] == "FONTE 70A":
                    return round(i['COLUNA5'], 2);
                elif self.option_selected_new == "FONTE 70A LITE" and i['PRODUTO'] == "FONTE 70A LITE":
                    return round(i['COLUNA5'], 2);
                elif self.option_selected_new == "FONTE 90 BOB" and i['PRODUTO'] == "FONTE 90 BOB":
                    return round(i['COLUNA5'], 2);
                elif self.option_selected_new == "FONTE 120 BOB" and i['PRODUTO'] == "FONTE 120 BOB":
                    return round(i['COLUNA5'], 2);
                elif self.option_selected_new == "FONTE 120A LITE" and i['PRODUTO'] == "FONTE 120A LITE":
                    return round(i['COLUNA5'], 2);
                elif self.option_selected_new == "FONTE 120A" and i['PRODUTO'] == "FONTE 120A":
                    return round(i['COLUNA5'], 2);
                elif self.option_selected_new == "FONTE 200 BOB" and i['PRODUTO'] == "FONTE 200 BOB":
                    return round(i['COLUNA5'], 2);
                elif self.option_selected_new == "FONTE 200A LITE" and i['PRODUTO'] == "FONTE 200A LITE":
                    return round(i['COLUNA5'], 2);
                elif self.option_selected_new == "FONTE 200 MONO" and i['PRODUTO'] == "FONTE 200 MONO":
                    return round(i['COLUNA5'], 2);
                elif self.option_selected_new == "FONTE 200A" and i['PRODUTO'] == "FONTE 200A":
                    return round(i['COLUNA5'], 2);
            for index, i in df.iterrows():
                if self.option_selected_new == "FONTE 40A" and i['PRODUTO'] == "FONTE 40A":
                    return round(i['COLUNA5'], 2);
                elif self.option_selected_new == "FONTE 60A" and i['PRODUTO'] == "FONTE 60A":
                    return round(i['COLUNA5'], 2);
                elif self.option_selected_new == "FONTE 60A LITE" and i['PRODUTO'] == "FONTE 60A LITE":
                    return round(i['COLUNA5'], 2);
                elif self.option_selected_new == "FONTE 70A" and i['PRODUTO'] == "FONTE 70A":
                    return round(i['COLUNA5'], 2);
                elif self.option_selected_new == "FONTE 70A LITE" and i['PRODUTO'] == "FONTE 70A LITE":
                    return round(i['COLUNA5'], 2);
                elif self.option_selected_new == "FONTE 90 BOB" and i['PRODUTO'] == "FONTE 90 BOB":
                    return round(i['COLUNA5'], 2);
                elif self.option_selected_new == "FONTE 120 BOB" and i['PRODUTO'] == "FONTE 120 BOB":
                    return round(i['COLUNA5'], 2);
                elif self.option_selected_new == "FONTE 120A LITE" and i['PRODUTO'] == "FONTE 120A LITE":
                    return round(i['COLUNA5'], 2);
                elif self.option_selected_new == "FONTE 120A" and i['PRODUTO'] == "FONTE 120A":
                    return round(i['COLUNA5'], 2);
                elif self.option_selected_new == "FONTE 200 BOB" and i['PRODUTO'] == "FONTE 200 BOB":
                    return round(i['COLUNA5'], 2);
                elif self.option_selected_new == "FONTE 200A LITE" and i['PRODUTO'] == "FONTE 200A LITE":
                    return round(i['COLUNA5'], 2);
                elif self.option_selected_new == "FONTE 200 MONO" and i['PRODUTO'] == "FONTE 200 MONO":
                    return round(i['COLUNA5'], 2);
                elif self.option_selected_new == "FONTE 200A" and i['PRODUTO'] == "FONTE 200A":
                    return round(i['COLUNA5'], 2);
        elif tipo == "Premium":
            for index, i in db.iterrows():
                if self.option_selected_new == "FONTE 40A" and i['PRODUTO'] == "FONTE 40A":
                    return round(i['COLUNA7'], 2);
                elif self.option_selected_new == "FONTE 60A" and i['PRODUTO'] == "FONTE 60A":
                    return round(i['COLUNA7'], 2);
                elif self.option_selected_new == "FONTE 60A LITE" and i['PRODUTO'] == "FONTE 60A LITE":
                    return round(i['COLUNA7'], 2);
                elif self.option_selected_new == "FONTE 70A" and i['PRODUTO'] == "FONTE 70A":
                    return round(i['COLUNA7'], 2);
                elif self.option_selected_new == "FONTE 70A LITE" and i['PRODUTO'] == "FONTE 70A LITE":
                    return round(i['COLUNA7'], 2);
                elif self.option_selected_new == "FONTE 90 BOB" and i['PRODUTO'] == "FONTE 90 BOB":
                    return round(i['COLUNA7'], 2);
                elif self.option_selected_new == "FONTE 120 BOB" and i['PRODUTO'] == "FONTE 120 BOB":
                    return round(i['COLUNA7'], 2);
                elif self.option_selected_new == "FONTE 120A LITE" and i['PRODUTO'] == "FONTE 120A LITE":
                    return round(i['COLUNA7'], 2);
                elif self.option_selected_new == "FONTE 120A" and i['PRODUTO'] == "FONTE 120A":
                    return round(i['COLUNA7'], 2);
                elif self.option_selected_new == "FONTE 200 BOB" and i['PRODUTO'] == "FONTE 200 BOB":
                    return round(i['COLUNA7'], 2);
                elif self.option_selected_new == "FONTE 200A LITE" and i['PRODUTO'] == "FONTE 200A LITE":
                    return round(i['COLUNA7'], 2);
                elif self.option_selected_new == "FONTE 200 MONO" and i['PRODUTO'] == "FONTE 200 MONO":
                    return round(i['COLUNA7'], 2);
                elif self.option_selected_new == "FONTE 200A" and i['PRODUTO'] == "FONTE 200A":
                    return round(i['COLUNA7'], 2);
            for index, i in df.iterrows():
                if self.option_selected_new == "FONTE 40A" and i['PRODUTO'] == "FONTE 40A":
                    return round(i['COLUNA7'], 2);
                elif self.option_selected_new == "FONTE 60A" and i['PRODUTO'] == "FONTE 60A":
                    return round(i['COLUNA7'], 2);
                elif self.option_selected_new == "FONTE 60A LITE" and i['PRODUTO'] == "FONTE 60A LITE":
                    return round(i['COLUNA7'], 2);
                elif self.option_selected_new == "FONTE 70A" and i['PRODUTO'] == "FONTE 70A":
                    return round(i['COLUNA7'], 2);
                elif self.option_selected_new == "FONTE 70A LITE" and i['PRODUTO'] == "FONTE 70A LITE":
                    return round(i['COLUNA7'], 2);
                elif self.option_selected_new == "FONTE 90 BOB" and i['PRODUTO'] == "FONTE 90 BOB":
                    return round(i['COLUNA7'], 2);
                elif self.option_selected_new == "FONTE 120 BOB" and i['PRODUTO'] == "FONTE 120 BOB":
                    return round(i['COLUNA7'], 2);
                elif self.option_selected_new == "FONTE 120A LITE" and i['PRODUTO'] == "FONTE 120A LITE":
                    return round(i['COLUNA7'], 2);
                elif self.option_selected_new == "FONTE 120A" and i['PRODUTO'] == "FONTE 120A":
                    return round(i['COLUNA7'], 2);
                elif self.option_selected_new == "FONTE 200 BOB" and i['PRODUTO'] == "FONTE 200 BOB":
                    return round(i['COLUNA7'], 2);
                elif self.option_selected_new == "FONTE 200A LITE" and i['PRODUTO'] == "FONTE 200A LITE":
                    return round(i['COLUNA7'], 2);
                elif self.option_selected_new == "FONTE 200 MONO" and i['PRODUTO'] == "FONTE 200 MONO":
                    return round(i['COLUNA7'], 2);
                elif self.option_selected_new == "FONTE 200A" and i['PRODUTO'] == "FONTE 200A":
                    return round(i['COLUNA7'], 2);
        elif tipo == "NA":
            for index, i in db.iterrows():
                if self.option_selected_new == "FONTE 40A" and i['PRODUTO'] == "FONTE 40A":
                    return round(i['COLUNA3'], 2);
                elif self.option_selected_new == "FONTE 60A" and i['PRODUTO'] == "FONTE 60A":
                    return round(i['COLUNA3'], 2);
                elif self.option_selected_new == "FONTE 60A LITE" and i['PRODUTO'] == "FONTE 60A LITE":
                    return round(i['COLUNA3'], 2);
                elif self.option_selected_new == "FONTE 70A" and i['PRODUTO'] == "FONTE 70A":
                    return round(i['COLUNA3'], 2);
                elif self.option_selected_new == "FONTE 70A LITE" and i['PRODUTO'] == "FONTE 70A LITE":
                    return round(i['COLUNA3'], 2);
                elif self.option_selected_new == "FONTE 90 BOB" and i['PRODUTO'] == "FONTE 90 BOB":
                    return round(i['COLUNA3'], 2);
                elif self.option_selected_new == "FONTE 120 BOB" and i['PRODUTO'] == "FONTE 120 BOB":
                    return round(i['COLUNA3'], 2);
                elif self.option_selected_new == "FONTE 120A LITE" and i['PRODUTO'] == "FONTE 120A LITE":
                    return round(i['COLUNA3'], 2);
                elif self.option_selected_new == "FONTE 120A" and i['PRODUTO'] == "FONTE 120A":
                    return round(i['COLUNA3'], 2);
                elif self.option_selected_new == "FONTE 200 BOB" and i['PRODUTO'] == "FONTE 200 BOB":
                    return round(i['COLUNA3'], 2);
                elif self.option_selected_new == "FONTE 200A LITE" and i['PRODUTO'] == "FONTE 200A LITE":
                    return round(i['COLUNA3'], 2);
                elif self.option_selected_new == "FONTE 200 MONO" and i['PRODUTO'] == "FONTE 200 MONO":
                    return round(i['COLUNA3'], 2);
                elif self.option_selected_new == "FONTE 200A" and i['PRODUTO'] == "FONTE 200A":
                    return round(i['COLUNA3'], 2);
            for index, i in df.iterrows():
                if self.option_selected_new == "FONTE 40A" and i['PRODUTO'] == "FONTE 40A":
                    return round(i['COLUNA3'], 2);
                elif self.option_selected_new == "FONTE 60A" and i['PRODUTO'] == "FONTE 60A":
                    return round(i['COLUNA3'], 2);
                elif self.option_selected_new == "FONTE 60A LITE" and i['PRODUTO'] == "FONTE 60A LITE":
                    return round(i['COLUNA3'], 2);
                elif self.option_selected_new == "FONTE 70A" and i['PRODUTO'] == "FONTE 70A":
                    return round(i['COLUNA3'], 2);
                elif self.option_selected_new == "FONTE 70A LITE" and i['PRODUTO'] == "FONTE 70A LITE":
                    return round(i['COLUNA3'], 2);
                elif self.option_selected_new == "FONTE 90 BOB" and i['PRODUTO'] == "FONTE 90 BOB":
                    return round(i['COLUNA3'], 2);
                elif self.option_selected_new == "FONTE 120 BOB" and i['PRODUTO'] == "FONTE 120 BOB":
                    return round(i['COLUNA3'], 2);
                elif self.option_selected_new == "FONTE 120A LITE" and i['PRODUTO'] == "FONTE 120A LITE":
                    return round(i['COLUNA3'], 2);
                elif self.option_selected_new == "FONTE 120A" and i['PRODUTO'] == "FONTE 120A":
                    return round(i['COLUNA3'], 2);
                elif self.option_selected_new == "FONTE 200 BOB" and i['PRODUTO'] == "FONTE 200 BOB":
                    return round(i['COLUNA3'], 2);
                elif self.option_selected_new == "FONTE 200A LITE" and i['PRODUTO'] == "FONTE 200A LITE":
                    return round(i['COLUNA3'], 2);
                elif self.option_selected_new == "FONTE 200 MONO" and i['PRODUTO'] == "FONTE 200 MONO":
                    return round(i['COLUNA3'], 2);
                elif self.option_selected_new == "FONTE 200A" and i['PRODUTO'] == "FONTE 200A":
                    return round(i['COLUNA3'], 2);

    def parse_location(self, response):
        name = response.meta['name']
        url = response.meta['url']
        new_price_float = response.meta['price']
        tipo = response.meta['tipo']
        cupom = response.meta['cupom']
        parcelado = self.get_price_previsto(tipo)
        loja = response.meta['loja']
        lugar = response.xpath('//*[@id="profile"]/div/div[2]/div[1]/div[3]/p/text()').get()


        doc.add_paragraph(f'Modelo: {self.option_selected_new}')
        doc.add_paragraph(f'URL: {url}')
        doc.add_paragraph(f'Nome: {name}')
        doc.add_paragraph(f'Preço: {new_price_float}')
        doc.add_paragraph(f'Preço Previsto: {parcelado}')
        doc.add_paragraph(f'Loja: {loja}')
        doc.add_paragraph(f'Tipo: {tipo}')
        doc.add_paragraph(f'Lugar: {lugar}')
        doc.add_paragraph(f'Cupom: {cupom}')
        doc.add_paragraph("--------------------------------------------------------------------")
        doc.add_paragraph('')
        
        yield {
            'url': url,
            'name': name,
            'price': new_price_float,
            'price_previsto': parcelado,
            'loja': loja,
            'tipo': tipo,
            'lugar': lugar
        }
        doc.save(fr"dados/{self.option_selected_new}.docx")

        
        
